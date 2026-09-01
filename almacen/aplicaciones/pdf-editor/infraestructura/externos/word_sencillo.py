# -*- coding: utf-8 -*-
"""Conversión a Word DE RESPALDO, para cuando la buena no puede.

La conversión normal (`pdf2docx`, en cliente_conversiones.py) reconstruye la
maquetación: párrafos, tablas, imágenes y su colocación. Cuando le sale bien, el
resultado es un Word casi igual que el PDF.

Pero hay documentos con los que no puede. pdf2docx analiza página por página y,
si una le da un error, la descarta y sigue; cuando descarta TODAS, el único aviso
que da es «No parsed pages», y el usuario se queda sin archivo (medido con un
documento del usuario el 30-jul-2026: respuesta 400 y nada que descargar).

Este módulo es la red de abajo: saca lo que se puede leer del PDF —el texto y las
tablas— y monta un .docx sencillo. **Pierde la maquetación**, pero es texto
editable de verdad y el usuario se lleva su documento. Más vale un Word sencillo
que un mensaje de error.

Se usa SOLO como respaldo: la primera opción sigue siendo pdf2docx.
"""
import io
import logging

logger = logging.getLogger(__name__)

# Debajo de esto una «tabla» detectada no es una tabla (una línea suelta, un borde).
MINIMO_FILAS = 2
MINIMO_COLUMNAS = 2
# Menos letras que esto (sin contar los rótulos «Página N») es un documento sin texto
# aprovechable: casi seguro un escaneo. Devolver un Word en blanco engañaría.
MINIMO_LETRAS = 20


class SinTextoAprovechable(Exception):
    """El PDF no tiene texto que pasar a Word (suele ser un escaneo sin OCR)."""


def construir(contenido_pdf):
    """Monta un .docx con el texto y las tablas del PDF. Devuelve los bytes.

    Cada página del PDF es una sección del documento, separada por un salto y con
    su número, para que quien lo abra sepa de dónde viene cada trozo.
    """
    import fitz
    from docx import Document
    from docx.shared import Pt

    documento = Document()
    documento.core_properties.title = 'Documento convertido desde PDF'
    letras = 0
    tablas_puestas = 0

    pdf = fitz.open(stream=contenido_pdf, filetype='pdf')
    try:
        for numero, pagina in enumerate(pdf, start=1):
            if numero > 1:
                documento.add_page_break()
            titulo = documento.add_paragraph()
            trozo = titulo.add_run('Página %d' % numero)
            trozo.bold = True
            trozo.font.size = Pt(9)

            zonas_de_tabla = []
            for tabla in _tablas_de(pagina):
                zonas_de_tabla.append(tabla['recuadro'])
                _escribir_tabla(documento, tabla['filas'])
                tablas_puestas += 1
                letras += sum(len(str(c or '')) for f in tabla['filas'] for c in f)

            # El texto que NO está dentro de una tabla: si no, saldría dos veces.
            for parrafo in _parrafos_fuera_de(pagina, zonas_de_tabla):
                documento.add_paragraph(parrafo)
                letras += len(parrafo)
    finally:
        pdf.close()

    if tablas_puestas == 0 and letras < MINIMO_LETRAS:
        # Un .docx con solo los rótulos de página no le sirve a nadie, y encima
        # parecería que la conversión salió bien.
        raise SinTextoAprovechable(
            'El documento no tiene texto que se pueda pasar a Word: parece un escaneo '
            '(una foto de cada hoja). Pásalo primero por Digitalizar (OCR) y vuelve a '
            'exportar.')

    memoria = io.BytesIO()
    documento.save(memoria)
    return memoria.getvalue()


def _tablas_de(pagina):
    """Tablas de la página, con su recuadro. Lista vacía si no hay o no se puede.

    `find_tables` es de PyMuPDF y no viene en versiones antiguas: si no está, se
    sigue sin tablas en vez de quedarse sin conversión.
    """
    try:
        encontradas = pagina.find_tables()
    except Exception as e:
        logger.warning('no se pudieron buscar tablas en la página %s: %s', pagina.number, e)
        return []

    tablas = []
    for tabla in getattr(encontradas, 'tables', []):
        try:
            filas = tabla.extract()
        except Exception:
            continue
        if len(filas) < MINIMO_FILAS or not filas or len(filas[0]) < MINIMO_COLUMNAS:
            continue
        tablas.append({'recuadro': tuple(tabla.bbox), 'filas': filas})
    return tablas


def _escribir_tabla(documento, filas):
    """Escribe las filas como una tabla de Word de verdad (editable)."""
    columnas = max(len(f) for f in filas)
    tabla = documento.add_table(rows=0, cols=columnas)
    try:
        tabla.style = 'Table Grid'          # con bordes; si la plantilla no lo trae, da igual
    except KeyError:
        pass
    for fila in filas:
        celdas = tabla.add_row().cells
        for i, valor in enumerate(fila[:columnas]):
            celdas[i].text = '' if valor is None else str(valor).strip()


def _parrafos_fuera_de(pagina, zonas):
    """Los párrafos de la página que no caen dentro de ninguna tabla."""
    parrafos = []
    try:
        bloques = pagina.get_text('blocks')
    except Exception as e:
        logger.warning('no se pudo leer el texto de la página %s: %s', pagina.number, e)
        return parrafos

    for bloque in bloques:
        x0, y0, x1, y1 = bloque[:4]
        texto = (bloque[4] or '').strip()
        if not texto:
            continue
        centro = ((x0 + x1) / 2, (y0 + y1) / 2)
        if any(_dentro(centro, z) for z in zonas):
            continue                        # ya salió en la tabla
        for linea in texto.split('\n'):
            linea = linea.strip()
            if linea:
                parrafos.append(linea)
    return parrafos


def _dentro(punto, recuadro):
    x, y = punto
    x0, y0, x1, y1 = recuadro
    return x0 <= x <= x1 and y0 <= y <= y1
