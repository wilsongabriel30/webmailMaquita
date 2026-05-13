"""Content Extractor — Full-text extraction from emails, PDFs, DOCX, XLSX.

Provides deep search capability for eDiscovery: extracts searchable text
from email body (plain + HTML) and common attachment formats.
"""
import email
import email.policy
import hashlib
import io
import logging
import os
import re
import tempfile
from html.parser import HTMLParser
from typing import Optional

logger = logging.getLogger("compliance.extractor")


class _HTMLTextExtractor(HTMLParser):
    """Strip HTML tags, keep text content."""
    def __init__(self):
        super().__init__()
        self._parts: list[str] = []
        self._skip = False

    def handle_starttag(self, tag, attrs):
        if tag in ("script", "style"):
            self._skip = True

    def handle_endtag(self, tag):
        if tag in ("script", "style"):
            self._skip = False
        if tag in ("br", "p", "div", "tr", "li", "h1", "h2", "h3", "h4", "h5", "h6"):
            self._parts.append("\n")

    def handle_data(self, data):
        if not self._skip:
            self._parts.append(data)

    def get_text(self) -> str:
        return " ".join("".join(self._parts).split())


def html_to_text(html: str) -> str:
    """Convert HTML to plain text."""
    try:
        parser = _HTMLTextExtractor()
        parser.feed(html)
        return parser.get_text()
    except Exception:
        # Fallback: regex strip
        return re.sub(r"<[^>]+>", " ", html)


def extract_pdf_text(data: bytes) -> str:
    """Extract text from PDF bytes."""
    try:
        import pdfplumber
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            parts = []
            for page in pdf.pages[:50]:  # Limit to 50 pages
                text = page.extract_text()
                if text:
                    parts.append(text)
            return "\n".join(parts)
    except Exception as e:
        logger.debug("PDF extraction failed: %s", e)
        return ""


def extract_docx_text(data: bytes) -> str:
    """Extract text from DOCX bytes."""
    try:
        from docx import Document
        doc = Document(io.BytesIO(data))
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text)
        return "\n".join(parts)
    except Exception as e:
        logger.debug("DOCX extraction failed: %s", e)
        return ""


def extract_xlsx_text(data: bytes) -> str:
    """Extract text from XLSX bytes."""
    try:
        from openpyxl import load_workbook
        wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
        parts = []
        for ws in wb.worksheets[:10]:  # Limit sheets
            for row in ws.iter_rows(max_row=500, values_only=True):
                row_text = " ".join(str(c) for c in row if c is not None)
                if row_text.strip():
                    parts.append(row_text)
        wb.close()
        return "\n".join(parts)
    except Exception as e:
        logger.debug("XLSX extraction failed: %s", e)
        return ""


def extract_csv_text(data: bytes) -> str:
    """Extract text from CSV bytes."""
    try:
        text = data.decode("utf-8", errors="replace")
        return text[:100000]  # Limit size
    except Exception:
        return ""


# Map MIME types and extensions to extractors
_EXTRACTORS = {
    "application/pdf": extract_pdf_text,
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": extract_docx_text,
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet": extract_xlsx_text,
    "application/msword": extract_docx_text,  # Best effort
    "application/vnd.ms-excel": extract_xlsx_text,
    "text/csv": extract_csv_text,
    "text/plain": lambda d: d.decode("utf-8", errors="replace")[:100000],
    "text/html": lambda d: html_to_text(d.decode("utf-8", errors="replace")),
}

_EXT_MAP = {
    ".pdf": extract_pdf_text,
    ".docx": extract_docx_text,
    ".xlsx": extract_xlsx_text,
    ".xls": extract_xlsx_text,
    ".doc": extract_docx_text,
    ".csv": extract_csv_text,
    ".txt": lambda d: d.decode("utf-8", errors="replace")[:100000],
    ".html": lambda d: html_to_text(d.decode("utf-8", errors="replace")),
    ".htm": lambda d: html_to_text(d.decode("utf-8", errors="replace")),
}


class EmailContent:
    """Parsed email content with body text and attachment texts."""

    def __init__(self):
        self.subject: str = ""
        self.sender: str = ""
        self.recipients: str = ""
        self.cc: str = ""
        self.bcc: str = ""
        self.date: str = ""
        self.message_id: str = ""
        self.body_text: str = ""
        self.body_html: str = ""
        self.body_plain: str = ""  # Combined plain text from body
        self.attachments: list[dict] = []  # [{name, mime_type, size, text}]
        self.hash_sha256: str = ""
        self.size_bytes: int = 0

    @property
    def full_text(self) -> str:
        """All searchable text combined: headers + body + attachments."""
        parts = [
            self.subject,
            self.sender,
            self.recipients,
            self.cc,
            self.body_plain,
        ]
        for att in self.attachments:
            if att.get("text"):
                parts.append(att["text"])
        return "\n".join(p for p in parts if p)

    def matches_keywords(self, keywords: list[str]) -> list[str]:
        """Check which keywords match in full text. Case-insensitive."""
        full = self.full_text.lower()
        return [kw for kw in keywords if kw.lower() in full]

    def get_snippet(self, keywords: list[str] = None, max_len: int = 500) -> str:
        """Get a text snippet, prioritizing keyword context."""
        text = self.body_plain or self.subject
        if not text:
            return ""

        if keywords:
            text_lower = text.lower()
            for kw in keywords:
                pos = text_lower.find(kw.lower())
                if pos >= 0:
                    start = max(0, pos - 100)
                    end = min(len(text), pos + len(kw) + 400)
                    return f"...{text[start:end]}..."

        return text[:max_len]


def parse_email_file(filepath: str) -> Optional[EmailContent]:
    """Parse a Maildir email file and extract all content."""
    try:
        with open(filepath, "rb") as f:
            raw_data = f.read()

        content = EmailContent()
        content.size_bytes = len(raw_data)
        content.hash_sha256 = hashlib.sha256(raw_data).hexdigest()

        msg = email.message_from_bytes(raw_data, policy=email.policy.default)

        # Headers
        content.subject = str(msg.get("Subject", ""))
        content.sender = str(msg.get("From", ""))
        content.recipients = str(msg.get("To", ""))
        content.cc = str(msg.get("Cc", ""))
        content.bcc = str(msg.get("Bcc", ""))
        content.date = str(msg.get("Date", ""))
        content.message_id = str(msg.get("Message-ID", ""))

        # Body and attachments
        body_texts = []
        body_htmls = []

        for part in msg.walk():
            ct = part.get_content_type()
            cd = str(part.get("Content-Disposition", ""))
            filename = part.get_filename()

            if ct == "multipart":
                continue

            try:
                payload = part.get_payload(decode=True)
                if not payload:
                    continue
            except Exception:
                continue

            if filename or "attachment" in cd:
                # This is an attachment
                att_info = {
                    "name": filename or "unnamed",
                    "mime_type": ct,
                    "size": len(payload),
                    "text": "",
                }

                # Try to extract text from attachment
                extractor = _EXTRACTORS.get(ct)
                if not extractor and filename:
                    ext = os.path.splitext(filename)[1].lower()
                    extractor = _EXT_MAP.get(ext)

                if extractor:
                    try:
                        att_info["text"] = extractor(payload)
                    except Exception as e:
                        logger.debug("Attachment extraction failed for %s: %s", filename, e)

                content.attachments.append(att_info)

            elif ct == "text/plain":
                charset = part.get_content_charset() or "utf-8"
                try:
                    body_texts.append(payload.decode(charset, errors="replace"))
                except Exception:
                    body_texts.append(payload.decode("utf-8", errors="replace"))

            elif ct == "text/html":
                charset = part.get_content_charset() or "utf-8"
                try:
                    html = payload.decode(charset, errors="replace")
                    body_htmls.append(html)
                    body_texts.append(html_to_text(html))
                except Exception:
                    pass

        content.body_text = "\n".join(body_texts)
        content.body_html = "\n".join(body_htmls)
        content.body_plain = content.body_text

        return content

    except Exception as e:
        logger.warning("Failed to parse email %s: %s", filepath, e)
        return None


def search_maildir(
    maildir_base: str,
    mailbox: str,
    keywords: list[str],
    folders: list[str] = None,
    date_from: str = None,
    date_to: str = None,
    search_body: bool = True,
    search_attachments: bool = True,
    max_results: int = 500,
) -> list[dict]:
    """
    Search a user's Maildir for emails matching keywords.

    Args:
        maildir_base: Base vmail path (e.g., /var/vmail)
        mailbox: Email address (e.g., user@maquita.org)
        keywords: List of search terms
        folders: Specific folders to search (None = all)
        date_from/date_to: Date filters (ISO format)
        search_body: Whether to search in email body
        search_attachments: Whether to search in attachment content
        max_results: Maximum results to return

    Returns:
        List of dicts with match info
    """
    import email.utils
    from datetime import datetime

    # Resolve maildir path
    parts = mailbox.split("@")
    if len(parts) != 2:
        return []

    user, domain = parts
    user_maildir = os.path.join(maildir_base, domain, user, "Maildir")

    if not os.path.isdir(user_maildir):
        logger.warning("Maildir not found: %s", user_maildir)
        return []

    # Parse date filters
    dt_from = None
    dt_to = None
    if date_from:
        try:
            dt_from = datetime.fromisoformat(date_from.replace("Z", "+00:00"))
        except Exception:
            pass
    if date_to:
        try:
            dt_to = datetime.fromisoformat(date_to.replace("Z", "+00:00"))
        except Exception:
            pass

    # Determine which folders to scan
    scan_dirs = []
    if folders:
        for folder in folders:
            folder_path = os.path.join(user_maildir, f".{folder}") if folder != "INBOX" else user_maildir
            if os.path.isdir(folder_path):
                scan_dirs.append((folder, folder_path))
    else:
        # Scan all folders
        # INBOX
        scan_dirs.append(("INBOX", user_maildir))
        # Subfolders (start with .)
        try:
            for entry in os.listdir(user_maildir):
                full = os.path.join(user_maildir, entry)
                if entry.startswith(".") and os.path.isdir(full) and entry not in (".dovecot", ".dovecot.lda-dupes.locks"):
                    folder_name = entry[1:]  # Remove leading dot
                    scan_dirs.append((folder_name, full))
        except Exception:
            pass

    results = []
    scanned = 0

    for folder_name, folder_path in scan_dirs:
        for subdir in ("cur", "new"):
            msg_dir = os.path.join(folder_path, subdir)
            if not os.path.isdir(msg_dir):
                continue

            try:
                files = os.listdir(msg_dir)
            except PermissionError:
                continue

            for fname in files:
                if len(results) >= max_results:
                    break

                filepath = os.path.join(msg_dir, fname)
                if not os.path.isfile(filepath):
                    continue

                scanned += 1

                # Date filter by file mtime (rough filter before parsing)
                if dt_from or dt_to:
                    try:
                        mtime = os.path.getmtime(filepath)
                        file_dt = datetime.fromtimestamp(mtime)
                        if dt_from and file_dt < dt_from.replace(tzinfo=None):
                            continue
                        if dt_to and file_dt > dt_to.replace(tzinfo=None):
                            continue
                    except Exception:
                        pass

                # Parse email
                content = parse_email_file(filepath)
                if not content:
                    continue

                # Build searchable text based on options
                search_text = content.subject + "\n" + content.sender + "\n" + content.recipients
                if search_body:
                    search_text += "\n" + content.body_plain
                if search_attachments:
                    for att in content.attachments:
                        if att.get("text"):
                            search_text += "\n" + att["text"]

                # Check keywords
                search_lower = search_text.lower()
                matched_kw = [kw for kw in keywords if kw.lower() in search_lower]

                if matched_kw:
                    # Extract UID from filename (Dovecot format: timestamp.xxx.host,S=size:2,flags)
                    uid = None
                    try:
                        uid = int(fname.split(".")[0])
                    except Exception:
                        pass

                    # Parse email date
                    sent_at = None
                    if content.date:
                        try:
                            parsed = email.utils.parsedate_to_datetime(content.date)
                            sent_at = parsed.isoformat()
                        except Exception:
                            pass

                    results.append({
                        "mailbox": mailbox,
                        "folder": folder_name,
                        "uid": uid,
                        "message_id": content.message_id,
                        "subject": content.subject,
                        "sender": content.sender,
                        "recipients": content.recipients,
                        "sent_at": sent_at,
                        "size_bytes": content.size_bytes,
                        "has_attachments": len(content.attachments) > 0,
                        "attachment_names": [a["name"] for a in content.attachments],
                        "matched_keywords": matched_kw,
                        "snippet": content.get_snippet(keywords),
                        "hash_sha256": content.hash_sha256,
                        "storage_path": filepath,
                        "search_scope": {
                            "body": search_body,
                            "attachments": search_attachments,
                        },
                    })

    logger.info(
        "eDiscovery search: mailbox=%s, keywords=%s, scanned=%d, matches=%d",
        mailbox, keywords, scanned, len(results),
    )
    return results
